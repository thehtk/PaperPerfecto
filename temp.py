import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
import helper
import re

st.set_page_config(
    page_title="Document Creator",
    page_icon="✍️",
)

def generate_word_document(title, author, abstract, keywords, body, references,authors_info):
    document = Document()
    
    #tittle
    para = document.add_paragraph()
    para.indentation_left = Inches(0.5)
    run = para.add_run(f"{title}")
    run.font.name = 'Times New Roman'
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run.font.size = Pt(24)
    run = para.runs[0]
    run.bold = True
    
    
    #author demo
    table = document.add_table(rows=1, cols=3)
    table.autofit = True  # Autofit the table to the content
    row_idx = 0
    col_idx = 0
    for author in authors_info:
        # Add author information to the table cell
        cell = table.cell(row_idx, col_idx)
        cell.text = author
        # Apply formatting to the cell text
        para = cell.paragraphs[0]
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para.indentation_left = Inches(0.5)
        run = para.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        # Move to the next column or row when the third column is filled
        col_idx += 1
        if col_idx == 3:
            col_idx = 0
            row_idx += 1
    
    
    #abstract
    para = document.add_paragraph()
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para.indentation_left = Inches(0.5)
    run = para.add_run("Abstract—")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)
    run.italic = True
    run.bold = True
    # Add the abstract text with regular formatting
    para.add_run('' + abstract).font.name = 'Times New Roman'
    para.runs[-1].font.size = Pt(9)
    para.runs[-1].bold = True
    
    
    
    #keywords
    para = document.add_paragraph()
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para.indentation_left = Inches(0.5)
    run = para.add_run(f"Keyword - {keywords}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = para.runs[0]
    run.italic = True
    run.bold = True
    
    
    """#body
    para = document.add_paragraph()
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para.indentation_left = Inches(0.5)
    run = para.add_run(f"Body - {body}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = para.runs[0] """
    
    
    #body demo
    section = document.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    pattern = r'<(section|subsection|body|equations)>(.*?)<\/tah>'
    matches = re.findall(pattern, body, re.DOTALL)

    for tag, content in matches:
        para = document.add_paragraph(content.strip())

    

        if tag == 'section':
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = para.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.bold = True

        elif tag == 'subsection':
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = para.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.italic = True
            

        elif tag == 'body':
            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            para.space_after = Inches(0.12)
            run = para.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

        elif tag == 'equations':
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = para.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.italic = True
                
    
    #demo ref
    for i, reference in enumerate(references, start=1):
        para = document.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        para.indentation_left = Inches(0.12)
        run = para.add_run(f"[{i}]  {reference}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8)
        
    
    
    return document

def main():
    st.markdown(
        """
        <style>
        body {
            background-color: #f0f2f6;
            padding: 20px;
        }
        </style>
        """,
        unsafe_allow_html=True
    )
    st.title("Document Creator")
    
    st.subheader("Title:")
    with st.expander("Enter the Title:",expanded=True):
        title = st.text_input("Enter the Title:",label_visibility="hidden")
    
    st.subheader("Author:")
    author=helper.get_author_info()
    
    st.subheader("Abstract:")
    with st.expander("Enter the Abstract:",expanded=True):
        abstract = st.text_area("Enter the Abstract:",label_visibility="hidden")
    
    st.subheader("Keywords")
    with st.expander("Enter the Keywords (comma-separated):",expanded=True):
        keywords = st.text_input("Enter Keywords (comma-separated):",label_visibility="hidden")
    
    st.subheader("Body:")
    with st.expander("Enter the Body:",expanded=True):
        body = st.text_area("Enter the Body:",label_visibility="hidden")
    
    st.subheader("Enter the References")
    references=helper.get_ref()

    if st.button("Generate Document"):
        document = generate_word_document(title, author, abstract, keywords, body, references,author)
        filename = f"{title.replace(' ', '_')}_document.docx"

        # Save the document to a BytesIO object
        from io import BytesIO
        buf = BytesIO()
        document.save(buf)
        buf.seek(0)
        st.balloons()

        # Provide a download link for the generated Word document
        st.download_button(label="Download Document", data=buf, file_name=filename)

if __name__ == "__main__":
    main()
